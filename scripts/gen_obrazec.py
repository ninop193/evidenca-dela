#!/usr/bin/env python3
"""Zgenerira izpolnljiv obrazec za evidenco delovnega časa v Word (.docx) in PDF.
Polja sledijo 18. členu ZEPDSV. Rezultat -> public/prenosi/."""
import os

OUT = os.path.join(os.path.dirname(__file__), "..", "public", "prenosi")
os.makedirs(OUT, exist_ok=True)

NASLOV = "Evidenca delovnega časa"
NAVODILA = [
    "Za vsakega delavca vodite ločen obrazec, en list na mesec. Vsak delovni dan izpolnite v svojo vrstico, sproti (ne za nazaj).",
    "Prihod / odhod: vpišite uro začetka in konca dela (npr. 8:00 in 16:00). Odmor: skupaj minut odmora med delom.",
    "Redne ure: skupaj opravljene redne ure v dnevu. Nadure: ure nad rednim delovnim časom.",
    "Posebni pogoji: ločeno vpišite ure nočnega, nedeljskega in prazničnega dela (npr. „nočno 2 / praznik 8“).",
    "Odsotnost: vrsto in ure odsotnosti z nadomestilom, ločeno v breme delodajalca ali drugih (npr. ZZZS), ter brez nadomestila (npr. „bolniška ZZZS 8“).",
    "Skupaj na dnu: tekoči seštevek ur za mesec (referenčno obdobje). Evidenco hranite kot listino trajne vrednosti.",
]
IDENT = [
    ("Delodajalec (naziv, sedež):", ""),
    ("Ime in priimek delavca:", ""),
    ("Delovno mesto:", ""),
    ("Mesec / leto:", ""),
    ("Polni / krajši delovni čas:", ""),
]
STOLPCI = [
    "Datum", "Prihod", "Odhod", "Odmor\n(min)", "Redne\nure", "Nad-\nure",
    "Posebni pogoji\n(nočno/ned./praz.)", "Odsotnost\n(vrsta + ure)", "Opomba",
]
VRSTIC = 31


# ---------------------------------------------------------------- WORD (.docx)
def build_docx(path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    # ležeče, ozki robovi
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.2))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    h = doc.add_paragraph()
    hr = h.add_run(NASLOV.upper())
    hr.bold = True
    hr.font.size = Pt(16)
    hr.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    sub = doc.add_paragraph()
    sr = sub.add_run("Obrazec po 18. členu Zakona o evidencah na področju dela in socialne varnosti (ZEPDSV)")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # identifikacijska polja (2 stolpca)
    it = doc.add_table(rows=0, cols=2)
    it.autofit = True
    for label, _ in IDENT:
        row = it.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = " " * 40  # prostor za vpis

    doc.add_paragraph()
    nl = doc.add_paragraph()
    nlr = nl.add_run("Navodila za izpolnjevanje")
    nlr.bold = True
    for n in NAVODILA:
        p = doc.add_paragraph(n, style="List Bullet")
        p.runs[0].font.size = Pt(8)

    doc.add_paragraph()

    # glavna tabela
    t = doc.add_table(rows=1, cols=len(STOLPCI))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, name in enumerate(STOLPCI):
        hdr[i].text = name
        for par in hdr[i].paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(8)
        # senčenje glave
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEFB")
        tcPr.append(shd)

    for _ in range(VRSTIC):
        cells = t.add_row().cells
        for c in cells:
            c.paragraphs[0].add_run("")
            c.paragraphs[0].runs and setattr(c.paragraphs[0].runs[0].font, "size", Pt(9))

    # seštevek
    sumrow = t.add_row().cells
    sumrow[0].text = "SKUPAJ"
    sumrow[0].paragraphs[0].runs[0].bold = True

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Evidenca je listina trajne vrednosti in se hrani trajno. Podpis delavca in delodajalca potrjuje pravilnost vpisanih ur."
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.rows[0].cells[0].text = "Podpis delavca: ______________________"
    sig.rows[0].cells[1].text = "Podpis delodajalca: ______________________"
    for cell in sig.rows[0].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.save(path)
    print("docx ->", path)


# ------------------------------------------------------------------- PDF
def build_pdf(path):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    )
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Registriraj pisavo s slovenskimi črkami (č/š/ž) — privzeta Helvetica jih nima.
    pdfmetrics.registerFont(TTFont("Slo", "/System/Library/Fonts/Supplemental/Arial.ttf"))
    pdfmetrics.registerFont(TTFont("Slo-Bold", "/System/Library/Fonts/Supplemental/Arial Bold.ttf"))
    pdfmetrics.registerFontFamily("Slo", normal="Slo", bold="Slo-Bold")

    doc = SimpleDocTemplate(
        path, pagesize=landscape(A4),
        leftMargin=10 * mm, rightMargin=10 * mm,
        topMargin=10 * mm, bottomMargin=10 * mm,
        title="Evidenca delovnega časa - obrazec",
    )
    styles = getSampleStyleSheet()
    brand = colors.HexColor("#1D4ED8")
    slate = colors.HexColor("#64748B")
    h1 = ParagraphStyle("h1", parent=styles["Title"], fontName="Slo-Bold", fontSize=17,
                        textColor=brand, alignment=TA_LEFT, spaceAfter=2)
    subs = ParagraphStyle("subs", parent=styles["Normal"], fontName="Slo", fontSize=8.5,
                          textColor=slate, spaceAfter=8, leading=11)
    lab = ParagraphStyle("lab", parent=styles["Normal"], fontName="Slo", fontSize=9, leading=16)
    nav = ParagraphStyle("nav", parent=styles["Normal"], fontName="Slo", fontSize=7.6, leading=10)
    navh = ParagraphStyle("navh", parent=styles["Normal"], fontName="Slo-Bold", fontSize=9,
                          spaceBefore=4, spaceAfter=2, textColor=colors.HexColor("#0F172A"))
    foot = ParagraphStyle("foot", parent=styles["Normal"], fontName="Slo", fontSize=7.6,
                          textColor=slate, leading=10, spaceBefore=6)

    story = [
        Paragraph("EVIDENCA DELOVNEGA ČASA", h1),
        Paragraph("Obrazec po 18. členu Zakona o evidencah na področju dela in socialne varnosti (ZEPDSV)", subs),
    ]

    # identifikacija
    id_rows = [[Paragraph("<b>" + l + "</b>", lab), ""] for l, _ in IDENT]
    id_tbl = Table(id_rows, colWidths=[62 * mm, 210 * mm])
    id_tbl.setStyle(TableStyle([
        ("LINEBELOW", (1, 0), (1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [id_tbl, Spacer(1, 6), Paragraph("<b>Navodila za izpolnjevanje</b>", navh)]
    for n in NAVODILA:
        story.append(Paragraph("•&nbsp;" + n, nav))
    story.append(Spacer(1, 8))

    # glavna tabela
    header = [c.replace("\n", "<br/>") for c in STOLPCI]
    hstyle = ParagraphStyle("hcell", fontName="Slo-Bold", fontSize=7.4, leading=8.5, alignment=1)
    data = [[Paragraph("<b>" + h + "</b>", hstyle) for h in header]]
    for _ in range(VRSTIC):
        data.append([""] * len(STOLPCI))
    data.append(["SKUPAJ"] + [""] * (len(STOLPCI) - 1))

    col_w = [18, 16, 16, 15, 15, 14, 46, 46, 45]
    col_w = [w * mm for w in col_w]
    tbl = Table(data, colWidths=col_w, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94A3B8")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEFB")),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F1F5F9")),
        ("FONTNAME", (0, 0), (-1, -1), "Slo"),
        ("FONTNAME", (0, -1), (0, -1), "Slo-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 1), (-1, -1), 4.5),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 4.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#FafBff")]),
    ]))
    story.append(tbl)
    story.append(Paragraph(
        "Evidenca je listina trajne vrednosti in se hrani trajno. Podpis delavca in delodajalca potrjuje pravilnost vpisanih ur.",
        foot))
    story.append(Spacer(1, 10))
    sig = Table([["Podpis delavca: ______________________",
                  "Podpis delodajalca: ______________________"]],
                colWidths=[136 * mm, 136 * mm])
    sig.setStyle(TableStyle([("FONTSIZE", (0, 0), (-1, -1), 9), ("FONTNAME", (0, 0), (-1, -1), "Slo")]))
    story.append(sig)

    doc.build(story)
    print("pdf  ->", path)


if __name__ == "__main__":
    build_docx(os.path.join(OUT, "evidenca-delovnega-casa-obrazec.docx"))
    build_pdf(os.path.join(OUT, "evidenca-delovnega-casa-obrazec.pdf"))
